import textract;    #Will alow .docx files to be read by python. installed using pip
import docx2txt;    #adds more options to word documents. Installed using pip
from unicodedata import *;  #ensures many unicode options are available. Built-in module
import docx;
from docx import Document;
import tkinter as tk;

#Two parameters: the first parameter is the name of the search query the second is for the file to scan.  
def cliffNoteBuilder(queryName, srcDoc):  #declare function
    try:
        window=tk.Tk();
        text = textract.process(srcDoc+'.docx', extension='docx', encoding="latin-1"); #latin-1 is the encoding needed to prevent the UnicodeError

        myString="";   #initialize string for later use to extract data from the text/word document.
        for x in text:  #pluck every element out of the textract object/Unit 5 Notes Copy.docx file.
            x=chr(x);   #Convert the ASCII characters into text
            myString+=x;    #Place every letter and word from the document inside the myString variable

        storeDoc=open(srcDoc+".txt","w", encoding="utf-8");    #Open a file object and put it in write mode

        print(myString, file=storeDoc)  #Print the myString variable inside of the .txt file.
        storeDoc.close()    #Close store.txt file in use process
        
        storyDocRead=open(srcDoc+".txt","r");  #Make the storeDoc.txt file readable
        storyLines=storyDocRead.readlines();    #Read all of the lines in 'storeDoc.txt' stored in variable: "storyDocRead"
        
        searchQuery=queryName;
        plusTabs="\t\t\t\t\t"+searchQuery;      #Tab the heading
        plusTabs=plusTabs[1:4]+plusTabs[5].upper()+plusTabs[6:];    #Format the Heading
        
        newDoc=Document();  #New Document object
        for x in storyLines:    #Pluck every element read from the storeDoc.txt/storyLines variable.
            if searchQuery.casefold() in x[:650].casefold() and len(x[:650])>300:     #If the string "Hoover" is in the first 174 characters of the line/element, perform below code block
                newDoc.add_heading(plusTabs);       #Add the heading from the search query
                newDoc.add_paragraph(x[:650]);      #Add a paragraph
                    #Add the lines containing "Hoover" into the 'hoover.doc'/'hooverFile' variable.
        
        newDoc.save(searchQuery+".docx");
    except PermissionError as p:
        print("\n\nWe have a problem! PLEASE make sure you close all word documents. \n");
        print(str(p)," \n\n \t\t\t\t");
    
#    for x in storyLines:    #Pluck every element read from the storeDoc.txt/storyLines variable.
 #       if searchQuery in x[:174].casefold():     #If the string "Hoover" is in the first 174 characters of the line/element, perform below code block
  #          print(searchQuery, file=depressFile);
   #         print(x, file=depressFile);      #Add the lines containing "Hoover" into the 'hoover.doc'/'hooverFile' variable.

cliffNoteBuilder("war", "Unit 7 World War 2 Notes");  #Makes a function call to the openWordDoc().
